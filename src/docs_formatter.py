# coding=utf-8
__author__ = 'gisly'

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_ORIENT
import re


HEADERS= [u'Предмет поиска (объект исследования, его составные части)',
          u'Страна выдачи (19), вид и номер охранного документа.\nКлассификационный индекс  МПК',
          u'Заявитель (71),  патентообладатель (73), страна. \
Номер заявки (21), дата приоритета (22), конвенционный приоритет (30), \
дата публикации заявки (43), патента (45)',
         u'Название изобретения и его сущность',
         u'Патенты-аналоги'
     ]


DOC_DELIM = '\n'


CHEM_ORDER = ['C',
'Si',
'Mn',
'Cr',
'Ni',
'Nb',
'W',
'Mo',
'V',
'Ti',
'Al',
'Co',
'Cu',
'S',
'P',
'N',
]

"""
creates a table for patents
TODO: try doing it using a document template
"""

def createPatentTable(patentList, docFileName):
    document = Document()
    for section in document.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
    table = document.add_table(rows=1, cols=5)
    fillInPatentHeader(table)
    fillInPatentListData(table, patentList)
    document.save(docFileName)

"""
creates a table containing chemical information
TODO: try doing it using a document template
"""    
    
def createChemicalTable(patentList, docFileName):
    document = Document()
    for section in document.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
    table = document.add_table(rows=1, cols=len(patentList) + 1)
    fillInChemicalData(table, patentList)
    document.save(docFileName)
    
   
"""
creates a table containing topic information
TODO: try doing it using a document template
"""      
    
def createTopicTable(patentList, docFileName):
    document = Document()
    for section in document.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
    table = document.add_table(rows=1, cols=len(patentList) + 1)
    fillInTopicData(table, patentList)
    document.save(docFileName)
    
    
    
def fillInPatentHeader(table):
    for i in range(0, len(HEADERS)):
        cell = table.cell(0, i)
        cell.text = HEADERS[i]
        
def fillInPatentListData(table, patentList):
    for i in range(0, len(patentList)):
        fillInPatentData(table, i, patentList[i])
        

        
def fillInPatentData(table, patentNum, patent):
    cells = table.add_row().cells
    for i in range(0, len(HEADERS)):
        cells[i].text = createColumnByIndex(i, patent).strip()
        
        
def fillInChemicalData(table, patentList):
    commonChemicalList = sortChemList(createCommonChemicalList(patentList))
    for chemical in commonChemicalList:
        newRow = table.add_row()
        newRow.cells[0].text = chemical
        for patentIndex, patent in enumerate(patentList):
            table.rows[0].cells[patentIndex + 1].text = createChemicalHeader(patent)
            chemText = getTextForChemicalByPatent(chemical, patent)
            newRow.cells[patentIndex + 1].text = chemText

"""sorts the chemical elements according to the standards
"""            
def sortChemList(chemList):
    return sorted(chemList, key = lambda chemElement:elementOrder(chemElement))  

def elementOrder(chemElement):
    if chemElement in CHEM_ORDER:
        return CHEM_ORDER.index(chemElement)
    return 100000000
     
"""formats chemical text"""            
def getTextForChemicalByPatent(chemical, patent):
    if not 'informationExtraction' in patent:
        return '-'
    chemInfo = patent['informationExtraction']['chemical']
    if chemical in chemInfo:
        lower = chemInfo[chemical][0]
        upper = chemInfo[chemical][1]
        if lower and upper:
            return str(lower['number'])+'-'+str(upper['number'])
        if lower:
            if lower['isIncl']:
                return '>='+str(lower['number'])
            return '>'+str(lower['number'])
        if upper:
            if upper['isIncl']:
                return '<='+str(upper['number'])
            return '<'+str(upper['number'])
    return '-'
    
    
def createChemicalHeader(patent):
    patentChemHeader = createCommonHeader(patent)
    informationExtractionAdditional = patent.get('informationExtractionAdditional')
    if informationExtractionAdditional:
        patentChemHeader += '\n' + informationExtractionAdditional
    return patentChemHeader


def createCommonHeader(patent):
    patentHeader = re.sub('[ABC]\d', '', patent['patentNumber']).strip() 
    owners =  patent.get('owners')
    if owners:
        patentHeader += '\n' + owners
    else:
        registrators =  patent.get('registrators')
        if registrators:
            patentHeader += '\n' + registrators
    
    publishedDate = patent.get('publishedDate')
    if publishedDate:
        patentHeader += '\n' + getYear(publishedDate)
    else:
        appPublishedDate = patent.get('appPublishedDate')
        if appPublishedDate:
            patentHeader += '\n' + getYear(appPublishedDate)
    return patentHeader


def fillInTopicData(table, patentList):
    new_row = table.add_row()
    for patentIndex, patent in enumerate(patentList):
        table.rows[0].cells[patentIndex + 1].text = createCommonHeader(patent)
        if 'informationExtraction' in patent and patent['informationExtraction']['topic'] is not None:
            new_row.cells[patentIndex + 1].text = patent['informationExtraction']['topic']
        else:
            new_row.cells[patentIndex + 1].text = ''
        
 
def getYear(yearDate):
    return yearDate.split('.')[-1] 
        
def createCommonChemicalList(patentList):
    commonChemicalList = set()
    for patent in patentList:
        if 'informationExtraction' in patent:
            commonChemicalList.update(patent['informationExtraction']['chemical'].keys())
    return commonChemicalList   
        
    
        
def createColumnByIndex(columnIndex, patent):    
    if columnIndex == 1:
        return createFirstColumn(patent)
    if columnIndex == 2:
        return createSecondColumn(patent) 
    if columnIndex == 3:
        return createThirdColumn(patent) 
    if columnIndex == 4:
        return createFourthColumn(patent) 
    return patent['originalId']


def createFirstColumn(patent):
    appData = []
    if patent.get('country'):
        appData += [patent['country']]
    appData.append(patent['type'])
    appData.append(patent['patentNumber'])
    appData +=patent['mpkClasses']
    return DOC_DELIM.join(appData)


def createSecondColumn(patent):
    regData = []
    if patent.get('registrators'):
        regData.append('(71) ' + patent['registrators'])
    if patent.get('owners'):
        regData.append('(73) ' + patent['owners'])
    regData.append('(21) ' +patent['appNumber'])
    regData.append('(22) ' +patent['appDate']) 
    
    
    priorityDates = patent.get('priorityDates')
    if priorityDates is not None:
        priorityStr = '(30)'
        for (priority, priorityDate) in priorityDates:
            priorityStr += ' '+priority+'\n' + priorityDate
        regData.append(priorityStr)
         
    if patent.get('appPublishedDate'):
        regData.append('(43) ' + patent['appPublishedDate'])
    if patent.get('publishedDate'):
        regData.append('(45) ' + patent['publishedDate'])
    if patent.get('pctNumber'):
        regData.append('(86) ' + patent['pctNumber'] + ' ' + patent['pctDate'])
    if patent.get('lawData'):
        regData.append('(84) ' + patent['lawData'])
    return DOC_DELIM.join(regData)


def createThirdColumn(patent):
    return DOC_DELIM.join([patent['title'], patent['abstractText']])

def createFourthColumn(patent):
    if patent.get('analogues'):
        return DOC_DELIM.join(patent['analogues'])
    return ''

        
    
    
       
    
        
